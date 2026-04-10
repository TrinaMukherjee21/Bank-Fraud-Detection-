import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('Detailed System Architecture & Technology Stack', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Overview
    doc.add_heading('1. System Overview', level=1)
    doc.add_paragraph(
        "The SecureGuard Fraud Detection System employs a robust, multi-layered architecture designed for high-throughput "
        "transaction processing. The architecture integrates cutting-edge machine learning libraries with a secure backend "
        "to provide real-time fraud probability scores and risk assessments."
    )

    # Architecture Diagram
    doc.add_heading('2. System Architecture Flowchart', level=1)
    
    # Path to the NEW detailed image
    image_path = r'C:\Users\mukhe\.gemini\antigravity\brain\11573e94-1145-4e87-84a8-d6fcc2910e96\system_architecture_detailed_tech_stack_bw_1774941823297.png'
    
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        caption = doc.add_paragraph('Figure 1: Detailed Vertical System Architecture with Modern Technology Stack')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Error: Detailed image not found! Please check the file path.]")

    # Technology Stack Details
    doc.add_heading('3. Technology Stack Breakdown', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technologies'
    hdr_cells[2].text = 'Purpose'

    tech_stack = [
        ('Frontend', 'HTML5, CSS3, JS, Chart.js, Bootstrap', 'Responsive analytics dashboard & real-time monitoring.'),
        ('Backend', 'Python 3.x, Flask, Gunicorn, Jinja2', 'REST API handling, core business logic, and routing.'),
        ('ML Pipeline', 'XGBoost, Scikit-learn, Pandas', 'Feature engineering and high-accuracy fraud classification.'),
        ('Security', 'Flask-Login, Bcrypt, PyJWT', 'Secure auth, password hashing, and session management.'),
        ('Data Store', 'SQLite 3, Joblib, CSV', 'Persistent storage for user data, logs, and model artifacts.')
    ]

    for layer, tech, purpose in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = purpose

    # Implementation Details
    doc.add_heading('4. Implementation Details', level=1)
    doc.add_paragraph(
        "The system utilizes an XGBoost Classifier trained on anonymized historical banking data. "
        "Real-time requests are validated, pre-processed using scikit-learn encoders, and then passed to the "
        "inference engine. Any transaction exceeding the risk threshold triggers an immediate security alert."
    )

    # Save
    save_path = r'c:\Bank Fraud Detection\reports\system_architecture_detailed.docx'
    doc.save(save_path)
    print(f"Detailed Report saved successfully to {save_path}")

if __name__ == "__main__":
    create_report()
