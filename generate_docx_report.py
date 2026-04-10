import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = docx.Document()

    # Title
    title = doc.add_heading('System Architecture for Real-Time Bank Fraud Detection', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "This document describes the high-level system architecture of the SecureGuard Fraud Detection System. "
        "The system is designed to provide real-time inference on transactional data using an XGBoost-based "
        "machine learning model integrated with a robust Flask backend."
    )

    # Architecture Diagram
    doc.add_heading('2. Architecture Flowchart', level=1)
    
    image_path = r'C:\Users\mukhe\.gemini\antigravity\brain\11573e94-1145-4e87-84a8-d6fcc2910e96\system_architecture_bw_minimalist_1774941656446.png'
    
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        caption = doc.add_paragraph('Figure 1: Vertical System Architecture Flowchart (Minimalist IEEE Style)')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Error: Image file not found. Please manually insert the architecture diagram.]")

    # Component Breakdown
    doc.add_heading('3. Component Breakdown', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Description'

    components = [
        ('Client Layer', 'Web-based interface for transaction input and real-time monitoring.'),
        ('Flask Web Server', 'Core backend handling routing, authentication, and session management.'),
        ('Preprocessing Unit', 'Transforms raw transactional data into model-compatible features.'),
        ('XGBoost Model', 'High-performance machine learning engine for fraud probability scoring.'),
        ('Risk Assessment', 'Logic-driven module for categorizing alerts and generating recommendations.'),
        ('Database Layer', 'Persistent storage for transactional logs and user accounts.')
    ]

    for comp, desc in components:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = desc

    # Save
    save_path = r'c:\Bank Fraud Detection\reports\system_architecture.docx'
    doc.save(save_path)
    print(f"Report saved successfully to {save_path}")

if __name__ == "__main__":
    create_report()
